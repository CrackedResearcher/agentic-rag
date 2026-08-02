"""Turning files of various formats into plain text.

Adding a format means writing one function and registering its extension in
``LOADERS``. Nothing else in the codebase needs to know the format exists.

Plain-text formats and HTML work with no extra dependencies. PDF and DOCX need
optional packages; if they are missing, the loader raises
:class:`MissingDependency` and ingest skips the file with a clear message
instead of crashing partway through a long run.
"""

from __future__ import annotations

import json
from collections.abc import Callable
from html.parser import HTMLParser
from pathlib import Path

# Files larger than this are skipped — usually binaries or data dumps that
# would blow up the embedding bill without adding retrievable knowledge.
MAX_FILE_BYTES = 20 * 1024 * 1024


class MissingDependency(RuntimeError):
    """A format was recognised but the package that reads it isn't installed."""

    def __init__(self, extension: str, package: str, extra: str) -> None:
        super().__init__(
            f"Reading {extension} files requires the '{package}' package. "
            f"Install it with: uv sync --extra {extra}"
        )


def load_text(path: Path) -> str:
    """Read a plain-text file, replacing anything that isn't valid UTF-8."""
    return path.read_text(encoding="utf-8", errors="replace")


def load_json(path: Path) -> str:
    """Pretty-print JSON so keys and values sit near each other in a chunk."""
    raw = load_text(path)
    try:
        return json.dumps(json.loads(raw), indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return raw  # malformed JSON is still searchable as text


class _TextExtractor(HTMLParser):
    """Collect visible text, skipping script and style contents."""

    SKIP = {"script", "style", "head", "meta", "link"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._skipping = 0

    def handle_starttag(self, tag: str, attrs: object) -> None:
        if tag in self.SKIP:
            self._skipping += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in self.SKIP and self._skipping:
            self._skipping -= 1

    def handle_data(self, data: str) -> None:
        if not self._skipping and data.strip():
            self.parts.append(data.strip())


def load_html(path: Path) -> str:
    """Strip tags from HTML using the standard library — no extra dependency."""
    parser = _TextExtractor()
    parser.feed(load_text(path))
    return "\n".join(parser.parts)


def load_pdf(path: Path) -> str:
    """Extract text from a PDF, page by page."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise MissingDependency(".pdf", "pypdf", "docs") from exc

    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page.strip() for page in pages if page.strip())


def load_docx(path: Path) -> str:
    """Extract paragraph and table text from a Word document."""
    try:
        import docx
    except ImportError as exc:
        raise MissingDependency(".docx", "python-docx", "docs") from exc

    document = docx.Document(str(path))
    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


#: Extension -> reader. The keys define which files ingest will pick up.
LOADERS: dict[str, Callable[[Path], str]] = {
    ".txt": load_text,
    ".md": load_text,
    ".markdown": load_text,
    ".rst": load_text,
    ".log": load_text,
    ".csv": load_text,
    ".tsv": load_text,
    ".yaml": load_text,
    ".yml": load_text,
    ".json": load_json,
    ".html": load_html,
    ".htm": load_html,
    ".pdf": load_pdf,
    ".docx": load_docx,
}

#: Formats needing an optional package, for help text and error messages.
OPTIONAL_FORMATS = {".pdf": "pypdf", ".docx": "python-docx"}


def supported_extensions() -> list[str]:
    return sorted(LOADERS)


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in LOADERS


def load(path: Path) -> str:
    """Read ``path`` as text using the loader registered for its extension.

    Raises:
        ValueError: the extension has no loader, or the file is too large.
        MissingDependency: the format needs a package that isn't installed.
    """
    loader = LOADERS.get(path.suffix.lower())
    if loader is None:
        raise ValueError(
            f"Unsupported file type '{path.suffix}'. Supported: {', '.join(supported_extensions())}"
        )

    size = path.stat().st_size
    if size > MAX_FILE_BYTES:
        raise ValueError(f"{path.name} is {size / 1_048_576:.1f} MB, over the 20 MB limit")

    return loader(path)
